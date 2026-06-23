from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
WORKSPACE_DIR = BASE_DIR.parent
TEMPLATE = WORKSPACE_DIR / "topic10_opening_report_template.docx"
OUTPUT = WORKSPACE_DIR / "topic10_entity_linking_agent_opening_report.docx"


TITLE = "实体链接与知识对齐智能体项目开题报告"
GROUP = "课题10实体链接与知识对齐智能体小组"

BACKGROUND = (
    "本课题面向数据治理与数据预处理工厂场景，研究“实体链接与知识对齐智能体”。在企业文档、"
    "业务报表、制度文件和知识库建设过程中，文本中会出现大量机构、人名、地点、产品、项目等"
    "实体指称。仅依靠实体识别只能找出名称，不能解决同名异指、简称别名、曾用名、知识库缺失"
    "和上下文消歧等问题，难以支撑后续检索、入库、统计分析和RAG知识库构建。\n"
    "本系统聚焦“识别之后”的链接与标准化：在已给定文本和mention的前提下，将指称映射到知识库"
    "中的标准实体ID和标准全称；当知识库不存在对应实体时输出NIL；对低置信或歧义结果给出复核"
    "标记；同时保留候选、分数、上下文关键词和决策日志，满足数据治理对可追溯、可复现、可审计"
    "的要求。\n"
    "与通用NER工具相比，本系统不重复做实体识别，而是重点解决实体消歧、别名标准化、NIL检测和"
    "简单共指回链问题。系统采用LangChain + LangGraph组织智能体流程，便于后续将规则召回替换为"
    "向量检索、图谱检索、重排序模型或大模型判别节点。"
)

GOALS = (
    "系统建设目标如下：\n"
    "1. 构建一个可独立部署、可被数据治理流水线调用的实体链接智能体服务，对外提供HTTP API接口。\n"
    "2. 支持输入文本、已识别实体指称mention和知识库，输出标准实体ID、标准全称、链接状态、置信度、候选列表、NIL判定和链接依据。\n"
    "3. 使用LangGraph编排实体链接流程，形成load_kb、generate_candidates、rerank_candidates、resolve_mentions、build_response、persist_trace等节点，保证流程清晰、可扩展、可追踪。\n"
    "4. 建立示例行业知识库和样例评测集，提供可复现的评测脚本，统计链接准确率、NIL precision/recall/F1、别名标准化召回率等指标。\n"
    "5. 支持trace_id贯穿全流程，落盘保存节点事件、决策日志和结果摘要，便于审计、回放与badcase分析。"
)

ARCHITECTURE = (
    "系统采用“接口层 + LangGraph编排层 + 链接能力层 + 知识库层 + 追踪评测层”的结构。接口层使用"
    "FastAPI暴露/v1/link、/v1/link/batch和/v1/traces/{trace_id}；编排层使用LangGraph StateGraph定义"
    "工作流，并用LangChain RunnableLambda封装每个节点；能力层包含候选召回、上下文重排、链接决策、"
    "NIL检测和简单共指回链；知识库层当前采用JSON实体库，后续可扩展到MySQL、Neo4j或向量数据库；"
    "追踪评测层负责保存trace和运行样例benchmark。\n"
    "核心工作流：START -> load_kb -> generate_candidates -> rerank_candidates -> resolve_mentions -> "
    "build_response -> persist_trace -> END。"
)

DIVISION = (
    "初步分工安排如下，可在确定小组成员后替换姓名：\n"
    "组长：负责总体方案设计、进度管理、接口规范确认和最终汇报材料整合。\n"
    "组员1：负责FastAPI接口、Pydantic数据模型、OpenAPI文档和服务部署脚本。\n"
    "组员2：负责知识库结构设计、实体别名维护、候选实体召回和样例知识库构建。\n"
    "组员3：负责LangChain + LangGraph工作流编排、节点事件记录、trace持久化和回放接口。\n"
    "组员4：负责评测数据集、evaluate.py评测脚本、指标统计、badcase分析和技术报告撰写。"
)

TECH_ROWS = [
    ("类别", "语言/框架/产品", "版本"),
    ("开发语言", "Python", "3.9+"),
    ("接口服务", "FastAPI + Uvicorn", "0.115.x / 0.30.x"),
    ("智能体编排", "LangChain + LangGraph", "0.3.x / 0.2.x"),
    ("数据校验", "Pydantic", "2.10.x"),
    ("知识库", "JSON实体库（可扩展MySQL/Neo4j/向量库）", "sample-energy-v1"),
    ("评测测试", "unittest + evaluate.py", "内置样例集"),
    ("接口文档", "OpenAPI / Swagger UI", "FastAPI内置"),
    ("部署方式", "venv / Docker", "本地或私有化部署"),
]


def set_paragraph_text(paragraph, text):
    paragraph.clear()
    lines = text.split("\n")
    run = paragraph.add_run(lines[0])
    for line in lines[1:]:
        run.add_break()
        run.add_text(line)


def format_run(run, size=10.5, bold=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold


def apply_document_font(document):
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0, 0, 0)

    headings = {"背景及意义", "系统建设目标", "技术架构", "小组分工"}
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        is_title = text == TITLE
        is_heading = text in headings
        if is_title:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            format_run(run, bold=True if is_title or is_heading else None)

    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        format_run(run, bold=True if row_index == 0 else None)


def fill_report():
    document = Document(TEMPLATE)

    replacements_by_index = {
        1: TITLE,
        13: f"小组：{GROUP}",
        14: "组长：待填写",
        15: "组员：待填写",
        16: "组员：待填写",
        17: "组员：待填写",
        20: BACKGROUND,
        22: GOALS,
        24: ARCHITECTURE,
        25: "主要技术选型如下表所示：",
        28: DIVISION,
    }

    for index, text in replacements_by_index.items():
        set_paragraph_text(document.paragraphs[index], text)

    if document.tables:
        table = document.tables[0]
        while len(table.rows) < len(TECH_ROWS):
            table.add_row()
        for row_index, row_values in enumerate(TECH_ROWS):
            for cell_index, value in enumerate(row_values):
                table.cell(row_index, cell_index).text = value
        for row_index in range(len(TECH_ROWS), len(table.rows)):
            for cell in table.rows[row_index].cells:
                cell.text = ""

    apply_document_font(document)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    fill_report()
